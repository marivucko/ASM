import math
import numpy as np
import matplotlib.pyplot as plt
import networkx as nx
import pandas as pd
from enum import Enum
from scipy.stats import pearsonr
import seaborn as sns
from matplotlib import pyplot
from selenium import webdriver
from selenium.common.exceptions import WebDriverException, NoSuchElementException
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
import time
from selenium.webdriver.common.by import By
import pickle
from datetime import datetime
from collections import Counter
from operator import itemgetter
import prettytable
from scipy import linalg
from docx import Document
from sklearn.cluster import SpectralClustering
from os.path import exists
from collections import defaultdict


class Type(Enum):
    SUBMISSION = 'submission'
    COMMENT = 'comment'


PATH = "C:\Program Files (x86)\chromedriver_win32\chromedriver_win32 (2)\chromedriver.exe"
alpha = 0.5
# date_to_compare = date.today()
date_to_compare = datetime.strptime('2009-1-1', '%Y-%m-%d')
options = Options()
options.add_argument('--no-sandbox')

global submissions, comments


def get_element_by_xpath(driver, xpath):
    try:
        return driver.find_element(By.XPATH, xpath).text
    except NoSuchElementException:
        return ''


def get_elements_by_xpath(driver, xpath):
    try:
        elements = driver.find_elements(By.XPATH, xpath)
        text = ''
        for element in elements:
            text = text + element.text + "\n"
        return text
    except NoSuchElementException:
        return ''


def read_from_csv(t):
    global submissions, comments
    data = []
    for i in range(0, 12):
        data.append(pd.read_csv(f'input_data/reddit2008/{t.value}s_2008_asm/csv-{i}.csv', low_memory=False, index_col=0, header=0))
    if t == Type.SUBMISSION:
        submissions = pd.concat(data, axis=0, ignore_index=True)
    elif t == Type.COMMENT:
        comments = pd.concat(data, axis=0, ignore_index=True)


def _1b():
    grupisano1 = submissions.loc[:, ['subreddit', 'subreddit_id', 'author']]
    grupisano2 = comments.loc[:, ['subreddit', 'subreddit_id', 'author']]
    gr = pd.concat([grupisano1, grupisano2])
    gr = gr[gr['author'] != '[deleted]'].groupby(['subreddit', 'subreddit_id'])['author'].agg(
        'nunique').reset_index().nlargest(10, 'author')
    print('a1\n', gr.to_string(index=False))


def _2():
    num_of_subreddit = submissions['subreddit_id'].nunique() + comments['subreddit_id'].nunique()
    grupisano1 = submissions.loc[:, ['subreddit', 'subreddit_id', 'author']]
    grupisano2 = comments.loc[:, ['subreddit', 'subreddit_id', 'author']]
    gr = pd.concat([grupisano1, grupisano2])
    num_of_authors = gr['author'].nunique()
    print('num_of_subreddit', num_of_subreddit)
    print('num_of_authors', num_of_authors)
    print('number of authors per subreddit', num_of_authors / num_of_subreddit)


def _3():
    print('---3---')
    grupisano1 = submissions.loc[:, ['subreddit', 'subreddit_id', 'author']]
    grupisano2 = comments.loc[:, ['subreddit', 'subreddit_id', 'author']]
    grupisano1 = grupisano1[grupisano1['author'] != '[deleted]'].groupby(['author'])['subreddit_id'].agg('nunique').reset_index().nlargest(10, 'subreddit_id')
    grupisano2 = grupisano2[grupisano2['author'] != '[deleted]'].groupby(['author'])['subreddit_id'].agg('nunique').reset_index().nlargest(10, 'subreddit_id')
    print('a3\n', grupisano1.rename(columns={'subreddit_id': 'Number of submissions by author'}).to_string(index=False))
    print('a1\n', grupisano2.rename(columns={'subreddit_id': 'Number of comments by author'}).to_string(index=False))


def _4():
    print('---4---')
    grupisano1 = submissions.loc[:, ['subreddit', 'subreddit_id', 'author']]
    grupisano2 = comments.loc[:, ['subreddit', 'subreddit_id', 'author']]
    gr = pd.concat([grupisano1, grupisano2])
    gr = gr[gr['author'] != '[deleted]'].groupby(['author'])['subreddit_id'].agg('nunique').reset_index().nlargest(10, 'subreddit_id')
    print('a3\n', gr.rename(columns={'subreddit_id': 'Number of subreddits on which the author is active'}).to_string(index=False))


def _5():
    print('---5---')
    grupisano1 = submissions.loc[:, ['subreddit', 'subreddit_id', 'author']]
    grupisano2 = comments.loc[:, ['subreddit', 'subreddit_id', 'author']]
    grupisano1 = grupisano1[grupisano1['author'] != '[deleted]'].groupby(['author'])['subreddit_id'].agg(
        'nunique').reset_index().rename(columns={'subreddit_id': 'Submissions'})
    grupisano2 = grupisano2[grupisano2['author'] != '[deleted]'].groupby(['author'])['subreddit_id'].agg(
        'nunique').reset_index().rename(columns={'subreddit_id': 'Comments'})
    gr = pd.merge(grupisano1, grupisano2, on='author', how='outer').fillna(0)
    gr['more'] = gr['Submissions'] + gr['Comments'] #pearsonr(gr['Submissions'], gr['Comments'])
    print('a3\n', gr)
    print(pearsonr(gr['Submissions'], gr['Comments']))
    print(gr.corr(method='pearson'))
    sns.scatterplot(x="Submissions", y="Comments", data=gr)
    pyplot.scatter(gr['Submissions'], gr['Comments'])
    pyplot.show()


def _6():
    print('---6---')
    a = submissions.nlargest(10, 'num_comments').fillna('').reset_index()
    s = Service(PATH)
    op = webdriver.ChromeOptions()
    driver = webdriver.Chrome(service=s, options=op)
    driver.set_page_load_timeout(30)
    for i in range(0, 10):
        time.sleep(0.5)
        row = a.iloc[[i], 1:]
        print(row.to_string(index=False))
        if row['over_18'].to_string(index=False) == 'False':
            try:
                print('@', a.iloc[i]['url'])
                driver.get(a.iloc[i]['url'])
                if a.iloc[i]['permalink'] in a.iloc[i]['url']:
                    print('---\n', get_element_by_xpath(driver, "//div[contains(@class, '_2SdHzo12ISmrC8H86TgSCp _29WrubtjAcKqzJSPdQqQ4h')]"))
                    print('------------------------------\n')
                else:
                    print('---\n', get_elements_by_xpath(driver, "//h1"))
                    print('\n', get_elements_by_xpath(driver, "//p"))
                    print('------------------------------\n')
            except WebDriverException as e:
                print('greska')
                print(e)
    # driver.close()
    # driver.quit()


def read_from_csv_clean_and_save_to_pickle_file():
    read_from_csv(Type.SUBMISSION)
    read_from_csv(Type.COMMENT)
    with open('reddit_data_cleaned/submissions_cleaned', 'wb') as file:
        pickle.dump(submissions, file)
    with open('reddit_data_cleaned/comments_cleaned', 'wb') as file:
        pickle.dump(comments, file)


def read_from_pickle_file():
    global submissions, comments
    submissions = pd.read_pickle(r'reddit_data_cleaned/submissions_cleaned')
    comments = pd.read_pickle(r'reddit_data_cleaned/comments_cleaned')


def make_subreddit_id__subreddit_dictionary():
    global submissions, comments
    read_from_pickle_file()
    data = pd.concat([submissions, comments])[['subreddit', 'subreddit_id']]
    data = data.drop_duplicates()
    subreddit_id__subreddit_dictionary = dict(zip(data.subreddit_id, data.subreddit))
    with open('reddit_data_cleaned/subreddit_id__subreddit_dictionary', 'wb') as file:
        pickle.dump(subreddit_id__subreddit_dictionary, file)


def get_subreddit_name_by_subreddit_id(subreddit_id):
    if not exists(r'reddit_data_cleaned/subreddit_id__subreddit_dictionary'):
        make_subreddit_id__subreddit_dictionary()
    subreddit_id__subreddit_dictionary = pd.read_pickle(r'reddit_data_cleaned/subreddit_id__subreddit_dictionary')
    return subreddit_id__subreddit_dictionary[subreddit_id]


def get_subreddit_id_by_subreddit_name(subreddit_name):
    if not exists(r'reddit_data_cleaned/subreddit_id__subreddit_dictionary'):
        make_subreddit_id__subreddit_dictionary()
    subreddit_id__subreddit_dictionary = pd.read_pickle(r'reddit_data_cleaned/subreddit_id__subreddit_dictionary')
    for subreddit_id, subreddit in subreddit_id__subreddit_dictionary.items():
        if subreddit == subreddit_name:
            return subreddit_id
    return None


def make_graph():
    G = nx.Graph()
    data = pd.concat([submissions, comments])[['subreddit', 'subreddit_id', 'author', 'created_utc']]
    # print(data)
    data = data.groupby(['subreddit', 'subreddit_id', 'author'])['created_utc'].agg('max').reset_index()
    # print(data)
    data = data[data.author != '[deleted]'].reset_index() # izbaceni cvorovi koji imaju samo deleted kao authora, mozda je jedan autor kreirao jednu obajvu i posle se obrisao
    # print(data)
    subreddits = data['subreddit_id'].unique()
    # print(subreddits, data['subreddit_id'].nunique())
    G.add_nodes_from(subreddits)
    all_authors = data['author'].unique()
    print(all_authors)
    l = 0
    for author in all_authors:
        #print('-', author)
        d = data.loc[data['author'] == author]
        #print(d, "\n")
        l = l + 1
        for i in range(0, len(d) - 1):
            for j in range(i + 1, len(d)):
                print('^', l, '/', len(all_authors), ' ', i, j)
                later_utc = max(d.iloc[i]['created_utc'], d.iloc[j]['created_utc'])
                later_date = datetime.fromtimestamp(later_utc)
                #print(d.iloc[i]['subreddit_id'], d.iloc[j]['subreddit_id'], later_utc)
                n = (date_to_compare.year - later_date.year) * 12 + date_to_compare.month - later_date.month
                if (d.iloc[i]['subreddit_id'], d.iloc[j]['subreddit_id']) in G.edges:
                    if G.edges[d.iloc[i]['subreddit_id'], d.iloc[j]['subreddit_id']]['weight'] < math.pow(alpha, n):
                        G.edges[d.iloc[i]['subreddit_id'], d.iloc[j]['subreddit_id']]['weight'] = math.pow(alpha, n)
                else:
                    G.add_edge(d.iloc[i]['subreddit_id'], d.iloc[j]['subreddit_id'], weight=math.pow(alpha, n))
    nx.write_gml(G, 'reddit_data_cleaned/graph.gml')


def make_s_net_t_graph(s_net):
    subreddit_name_list = ['reddit.com', 'pics', 'worldnews', 'programming', 'business', 'politics', 'obama', 'science',
                           'technology', 'WTF', 'AskReddit', 'netsec', 'philosophy', 'videos', 'offbeat', 'funny',
                           'entertainment', 'linux', 'geek', 'gaming', 'comics', 'gadgets', 'nsfw', 'news',
                           'environment', 'atheism', 'canada', 'math', 'Economics', 'scifi', 'bestof', 'cogsci', 'joel',
                           'Health', 'guns', 'photography', 'software', 'history', 'ideas']
    subreddit_id_list = []
    for subreddit_name in subreddit_name_list:
        subreddit_id_from_subreddit_name = get_subreddit_id_by_subreddit_name(subreddit_name)
        if subreddit_id_from_subreddit_name is not None:
            subreddit_id_list.append(subreddit_id_from_subreddit_name)
    print(subreddit_id_list)
    s_net_t = s_net.subgraph(subreddit_id_list)
    nx.write_gml(s_net_t, 'graphs/s_net_t.gml')
    return s_net_t


def read_s_net_t_graph(s_net):
    if exists(r'graphs/s_net_t.gml'):
        s_net_t = nx.read_gml('graphs/s_net_t.gml')
    else:
        s_net_t = make_s_net_t_graph(s_net)
    return s_net_t


def analise_s_net_t(s_net):
    s_net_t = read_s_net_t_graph(s_net)
    print(nx.info(s_net_t))
    # s_net_7(s_net_t)
    # s_net_8(s_net_t)
    # s_net_9(s_net_t)
    # s_net_10(s_net_t)
    # s_net_12_13(s_net_t)
    # s_net_14(s_net_t)
    # s_net_15(s_net_t)
    # s_net_16(s_net_t)
    # s_net_17(s_net_t)
    # s_net_18(s_net_t)
    # s_net_19(s_net_t)
    # spectral(s_net_t)


def make_user_net():
    read_from_pickle_file()
    s = submissions[submissions.author != '[deleted]'][['id', 'author', 'created_utc']].drop_duplicates().reset_index()
    s_id = s.id
    s_author_name = dict(zip(s_id, s.author))
    s_created_utc = dict(zip(s_id, s.created_utc))
    c = comments[comments.author != '[deleted]'][['id', 'parent_id', 'author', 'created_utc']].drop_duplicates().reset_index()
    c_id = c.id
    c_parent_ids = dict(zip(c_id, c.parent_id))
    c_author_names = dict(zip(c_id, c.author))
    c_created_utcs = dict(zip(c_id, c.created_utc))

    user_net = nx.MultiGraph()
    for i in range(0, len(c_id)):
        c_id_curr = c_id[i]
        c_parent_id = c_parent_ids[c_id_curr]
        c_author_name = c_author_names[c_id_curr]
        c_created_utc = c_created_utcs[c_id_curr]
        parent_type = c_parent_id[:2]
        parent_id = c_parent_id[3:]
        if parent_type == 't1':
            print('&1')
            # parent_author_id = c_author_name.get(parent_id)
            parent_author_name = c_author_names.get(parent_id)
            parent_author_created_utc = c_created_utcs.get(parent_id)
            p_type = Type.COMMENT.value
        else:
            print('&3')
            parent_author_name = s_author_name.get(parent_id)
            parent_author_created_utc = s_created_utc.get(parent_id)
            p_type = Type.SUBMISSION.value
        print(i, p_type, c_id_curr, c_parent_id, parent_type, parent_id, '*', c_author_name, parent_author_name)
        if parent_author_name is not None:
            # print(int(c_created_utc[c_id_curr]))
            # print(parent_author_name)
            # print(parent_author_created_utc)
            later_utc = max(c_created_utc, parent_author_created_utc)
            later_date = datetime.fromtimestamp(later_utc)
            n = (date_to_compare.year - later_date.year) * 12 + date_to_compare.month - later_date.month
            if (c_author_name, parent_author_name) in user_net.edges:
                if user_net.edges[c_author_name, parent_author_name]['weight'] < math.pow(alpha, n):
                    user_net.edges[c_author_name, parent_author_name]['weight'] = math.pow(alpha, n)
            else:
                user_net.add_edge(c_author_name, parent_author_name, weight=math.pow(alpha, n), key=p_type)
    print(nx.info(user_net))
    print(user_net.edges.data())
    nx.write_gml(user_net, 'graphs/user_net_with_self_loops.gml')





    # dd = defaultdict(list)
    # for d in (c_dict1, c_dict2, c_dict3):  # you can list as many input dicts as you want here
    #     for key, value in d.items():
    #         print(value)
    #         dd[key].append(value)
    # print(dd)



def read_graph_from_file():
    graph = nx.read_gml('reddit_data_cleaned/graph.gml')
    print(graph.number_of_nodes())
    print(graph.number_of_edges())
    return graph


def get_giant_component(graph):
    gcc = sorted(nx.connected_components(graph), key=len, reverse=True)
    dominant_component = graph.subgraph(gcc[0])
    return dominant_component


def s_net_7(graph):
    print(f'Gustina SNet mreže je { nx.density(graph) }.')


def s_net_8(graph):
    try:
        print(f'Prosečna distanca u SNet mrezi je { nx.average_shortest_path_length(graph) } , dok je dijametar { nx.diameter(graph) }.')
    except nx.exception.NetworkXError:
        print('Mreža nije povezana, stoga su prosečna dužina i dijametar mreže beskonačni.')


def s_net_9(graph):
    components = list(nx.connected_components(graph))
    n = 0
    x = []
    for i in range(0, len(components)):
        x.append(len(components[i]))
        if len(components[i]) == 1:
            n = n + 1
    x = sorted(x, reverse=True)
    print(f'Broj povezanih komponenti mreže je { len(components)}.\nVeličina komponenti je redom { x }.')
    print(f'Broj cvorova koji su nepovezani je { n }.')
    if x[0] / graph.number_of_nodes() > 0.5:
        print(f'Postoji gigantska komponenta mreže jer postoji jezgro veličine { x[0] } čvorova i predstavlja najveći deo mreže tj. { round(x[0] / graph.number_of_nodes() * 100, 2) } % mreze.')
    else:
        print('Ne postoji gigantska komponenta mreže jer veličina najveće komponente ne predstavlja i najveći deo mreže.')


def s_net_10(graph):
    print(f'Prosečni koeficijent klasterizacije je { nx.average_clustering(graph) }.')
    number_of_triangles = sum(nx.triangles(graph).values()) / 3
    n = graph.number_of_nodes()
    max_number_of_triangles = n * (n - 1) * (n - 2) / 6
    global_clustering_coefficient = number_of_triangles / max_number_of_triangles
    print(f'Globalni koeficijent klasterizacije je { global_clustering_coefficient }.')
    print(f'Lokalni koeficijenti klasterizacije čvorova { nx.clustering(graph) }.')

    # dot graph of local clustering coefficient distribution
    ids, lcc = zip(*nx.clustering(graph).items())
    lcc_counts = Counter(lcc)
    x1, y1 = zip(*lcc_counts.items())
    plt.scatter(x1, y1, marker='.')
    plt.show()

    # histogram graph of local clustering coefficient distribution
    # lcc = list(nx.clustering(graph).values())
    # for i in range(0, len(lcc)):
    #     print(i, lcc[i])
    # (ax1, ax2) = plt.subplots(ncols=1)
    # ax2.hist(lcc, bins=10)
    # ax2.set_xlabel('Local clustering coefficient')
    # ax2.set_ylabel('Frequency')
    # plt.tight_layout()
    # plt.show()

    # Erdos-Reny
    n = graph.number_of_nodes()
    m = graph.number_of_edges()
    H = nx.erdos_renyi_graph(n, 2 * float(m) / (n * (n - 1)))
    ids, lcc = zip(*nx.clustering(H).items())
    lcc_counts = Counter(lcc)
    x2, y2 = zip(*lcc_counts.items())
    plt.scatter(x2, y2, marker='.')
    plt.show()


def s_net_12_13(graph):
    r = nx.degree_assortativity_coefficient(graph)
    print(f'Koeficijent asortativnosti je { r }.')
    # da li su cvorovi sa visokim stepenom povezani sa cvorovima sa visokim stepenom
    ids, num_of_neighbours = zip(*sorted(graph.degree, key=lambda x: x[1], reverse=True))
    num_of_nodes_with_highest_degree = 100
    graph1 = graph.subgraph(ids[:num_of_nodes_with_highest_degree])
    labels = {}
    ids1, num_of_neighbours1 = zip(*sorted(graph1.degree, key=lambda x: x[1], reverse=True))
    for i in range(0, num_of_nodes_with_highest_degree):
        labels[ids1[i]] = f'Degree = { num_of_neighbours1[i] }'
    nx.draw(graph1, labels=labels, with_labels=True, font_color="whitesmoke")
    plt.show()


def plot_degree_distribution(graph):
    _, degree_list = zip(*graph.degree())
    degree_counts = Counter(degree_list)
    x, y = zip(*degree_counts.items())
    plt.xlabel('degree')
    plt.ylabel('frequency')
    plt.xscale('log')
    plt.yscale('log')
    plt.xlim(1, max(x))
    plt.ylim(1, max(y))
    plt.figure(1)
    plt.scatter(x, y, marker='.')
    plt.show()


def s_net_14(graph):
    plot_degree_distribution(graph)
    n = graph.number_of_nodes()
    m = graph.number_of_edges()
    H = nx.erdos_renyi_graph(n, 2 * float(m) / (n * (n - 1)))
    plot_degree_distribution(H)


def print_top_n_from_dictionary(dictionary, n):
    sorted_dictionary = dict(sorted(dictionary.items(), key=itemgetter(1), reverse=True)[:10])
    table = prettytable.PrettyTable(["Čvor", "Vrednost"])
    for item in sorted_dictionary:
        table.add_row([repr(item).strip("'"), sorted_dictionary[item]])
    print(table)
    print("\n")


def s_net_15(graph):
    # kod neusmerenog grafa hub i authorities su isti
    hubs, authorities = nx.hits(graph)
    print_top_n_from_dictionary(hubs, 10)
    #print_top_n_from_dictionary(authorities, 10)

    # proveravamo da li su hubs i authorities zapravo u jezgru i poveznai sa covorovima visokog stepena a i oni sami imaju visok stepen
    ids, num_of_neighbours = zip(*sorted(graph.degree, key=lambda x: x[1], reverse=True))
    max_num_of_neighbours = num_of_neighbours[0]
    i = 0
    while (num_of_neighbours[i] == max_num_of_neighbours) or (i < 10):
        print(i, '.', ids[i], num_of_neighbours[i])
        i = i + 1
    print('\n')
    # ispisuje se stepen cvorova koji su habovi
    sorted_dictionary = dict(sorted(hubs.items(), key=itemgetter(1), reverse=True)[:10])
    d = graph.degree
    for item in sorted_dictionary:
        print(repr(item), ":", sorted_dictionary[item], d[item])

    table = prettytable.PrettyTable(["Čvor", "Hub/Auth Vrednost", "Stepen čvor", "10 najvećih stepeni čvorova suseda zadatom čvoru"])
    for item in sorted_dictionary:
        node = repr(item).strip("'")
        degrees_of_node_neighbours = get_degrees_of_node_neighbours(graph, node)
        table.add_row([node, sorted_dictionary[item], d[item], degrees_of_node_neighbours])
    print(table, "\n")


def s_net_16(graph):
    print('10 čvorova sa najvećom vrednošću centralnosti po stepenu:')
    print_top_n_from_dictionary(nx.degree_centrality(graph), 10)
    print('10 čvorova sa najvećom vrednošću centralnosti po bliskosti:')
    print_top_n_from_dictionary(nx.closeness_centrality(graph), 10)
    print('10 čvorova sa najvećom vrednošću relacione centralnosti:')
    print_top_n_from_dictionary(nx.betweenness_centrality(graph), 10)


def get_degrees_of_node_neighbours(graph, node):
    ids, degrees = zip(*sorted(graph.degree, key=lambda x: x[1], reverse=True))
    neighbours = list(graph.neighbors(node))
    i = 0
    j = 0
    degrees_of_node_neighbours = []
    while j < 10:
        if ids[i] in neighbours:
            j = j + 1
            degrees_of_node_neighbours.append(degrees[i])
        i = i + 1
    return degrees_of_node_neighbours


def s_net_17(graph):
    print('10 čvorova sa najvećom vrednošću centralnosti po sopstvenom vektoru:')
    sorted_dictionary = dict(sorted(nx.eigenvector_centrality(graph, weight='weight').items(), key=itemgetter(1), reverse=True)[:10])
    table = prettytable.PrettyTable(["Čvor", "Vrednost", "10 najvećih stepeni čvorova suseda zadatom čvoru"])
    for item in sorted_dictionary:
        node = repr(item).strip("'")
        degrees_of_node_neighbours = get_degrees_of_node_neighbours(graph, node)
        table.add_row([node, sorted_dictionary[item], degrees_of_node_neighbours])
    print(table)
    print("\n")


def print_katz(graph, alfa, beta):
    beta_for_reddit_com = beta['t5_6']
    print(f'Katzova centralnost za alfa = { alfa } i beta za subredit reddit.com = { beta_for_reddit_com }, dok je za ostale 1')
    katz_dict = nx.katz_centrality(graph, alpha=alfa, beta=beta, tol=1e-06, max_iter=5000, nstart=None, normalized=True, weight='weight')
    print_top_n_from_dictionary(katz_dict, 10)


def s_net_18(graph):
    lambda_max = max(nx.adjacency_spectrum(graph))
    alfa_max = (1 / lambda_max)
    print('Alfa max:', alfa_max)
    alfa_max = (alfa_max * 0.99).real
    alfa_step = alfa_max / 10
    nodes = graph.nodes()
    beta_nodes_dictionary = dict.fromkeys(nodes, 1)
    alfa = 0
    while alfa < alfa_max:
        for i in range(0, 5):
            beta_nodes_dictionary['t5_6'] = 5 * i + 1
            print_katz(graph, alfa, beta_nodes_dictionary)
        alfa = alfa + alfa_step


def s_net_19(graph):
    degree_centrality = nx.degree_centrality(graph)
    closeness_centrality = nx.closeness_centrality(graph)
    betweenness_centrality = nx.betweenness_centrality(graph)
    eigenvector_centrality = nx.eigenvector_centrality(graph)
    composite_rank = dict.fromkeys(graph.nodes(), 0)
    for item in composite_rank:
        composite_rank[item] = degree_centrality[item] * closeness_centrality[item] * betweenness_centrality[item] * eigenvector_centrality[item]
    print('10 čvorova sa najvećom vrednošću kopozitnog ranga:')
    print_top_n_from_dictionary(composite_rank, 10)


def s_net_16_17_19(graph):
    degree_centrality = nx.degree_centrality(graph)
    closeness_centrality = nx.closeness_centrality(graph)
    betweenness_centrality = nx.betweenness_centrality(graph)
    eigenvector_centrality = nx.eigenvector_centrality(graph)
    composite_rank = dict.fromkeys(graph.nodes(), 0)
    table = prettytable.PrettyTable(["Čvor", "DC", "CC", "BC", "EVC", "Composite rank"])
    for item in composite_rank:
        composite_rank[item] = degree_centrality[item] * closeness_centrality[item] * betweenness_centrality[item] * eigenvector_centrality[item]
    composite_rank = dict(sorted(composite_rank.items(), key=itemgetter(1), reverse=True)[:10])
    for item in composite_rank:
        table.add_row([item, degree_centrality[item], closeness_centrality[item], betweenness_centrality[item], eigenvector_centrality[item], composite_rank[item]])
    print(table)
    print("\n")


def plot_scatter_plot(x_data, y_data, x_label, y_label, title):
    fig, ax = plt.subplots()
    ax.scatter(x_data, y_data)
    ax.set_ylabel(y_label, fontsize=15)
    ax.set_xlabel(x_label, fontsize=15)
    ax.set_title(title)
    plt.savefig(f'spectral/{title}.svg')
    plt.show()


def in_range(value, target, tolerated_error):
    return abs(value-target) < tolerated_error


def get_components3(L, component_number):
    evals, evecs = linalg.eig(L) # each column of evecs is an eigenvector\n",
    # take only real component of each eigenvalue, because eigenvalues of L are guaranteed to be real \n",
    evals = evals.real
    # take only real component of each eigenvector, because eigenvectors of L are guaranteed to be real \n",
    evecs_temp = []
    for e in evecs:
        v = e.real
        evecs_temp.append(v)
    evecs = evecs_temp
    evecs_horizontal_container = np.array(evecs).T #each row is an eigenvector\n",
    zipped = list(zip(evals, evecs_horizontal_container)) # pair eigenvalues and eigenvectors into eigenpairs\n",
    zipped.sort(key=lambda tup: tup[0]) # sort eigenpairs by eigenvalue\n",
    evals, evecs = zip(*zipped) # evecs is now horizontal, but ordered by paired eigenvalue size\n",
    eigenspace0_basis = evecs[:component_number] # take only the first component_number eigenvectors \n",
    #eigenspace0_basis = filter_matrix(eigenspace0_basis, math.pow(10, -15)) # filter values close to 0\n",
    return find_ccs2(eigenspace0_basis, len(eigenspace0_basis[0]))


def find_ccs2(eigenvectors_holder, n):
    components = []
    already_solved = []
    for current in range(0, n):
        if current not in already_solved:
            component = [current]
            for running in range (current+1, n):
                if running not in already_solved:
                    b = True
                    for tracer in range(0, len(eigenvectors_holder)):
                        eigenvector = eigenvectors_holder[tracer]
                        if not in_range(eigenvector[current], eigenvector[running], math.pow(10,-12)):
                            b = False
                    if b:
                        component.append(running)
                        already_solved.append(running)
            already_solved.append(current)
            components.append(component)
    return components


def grapf_laplasijan_s_net_20(graph):
    L = nx.laplacian_matrix(graph).toarray()
    eigenvalues = linalg.eigvals(L) # TODO: change to .eigenvalsh because we know the matrix is symmetric
    eigenvalues.sort()
    enumerator = np.array(range(1, len(eigenvalues)+1))
    df_eig = pd.DataFrame(list(zip(enumerator, eigenvalues)))
    # write eigenvalue table
    df_eig30 = df_eig[:30]
    df_eig30.columns = ['k', 'lambda_k']
    df_eig30 = df_eig30.astype({'k': 'int32', 'lambda_k':'float'})
    print(df_eig30)
    doc = Document()
    #doc = docx.Document('tabela_eig.docx')
    t = doc.add_table(df_eig30.shape[0]+1, df_eig30.shape[1])
    t.cell(0,0).text = r'$k$'
    t.cell(0,1).text = r'$lambda_k$'
    for i in range(df_eig30.shape[0]):
        for j in range(df_eig30.shape[-1]):
            t.cell(i+1,j).text = str(df_eig30.values[i,j])
    doc.save('spectral/tabela_eig.docx')
    plot_scatter_plot(enumerator, eigenvalues, r'$k$', r'$lambda_k$', 'Ceo spektar graf laplasijana')
    # prvih 30 sopstvenih vrednosti,
    # df_eig_30 = df_eig[:30]
    # u snetu krecemo od 800 jer oko tog broja krece nagli skok, videti da se ovde gleda samo gigantska komponenta pa ce biti i manje klastera
    df_eig_30 = df_eig[810:850]
    plot_scatter_plot(df_eig_30.iloc[:,0], df_eig_30.iloc[:,1], r'$k$', r'$lambda_k$', 'Prvih 30 sopstvenih vrednosti graf laplasijana')


def spectral(graph):
    dominant_component = get_giant_component(graph)
    for k in range(2, 30):
        clustering = SpectralClustering(n_clusters=k, assign_labels="discretize", affinity="precomputed").fit(nx.adjacency_matrix(dominant_component))
        colors = clustering.labels_
        c_string = []
        for c in colors:
            c_string.append(str(c))
        G1 = nx.Graph()
        for c, label in zip(c_string, dominant_component.nodes()):
            G1.add_node(label, color=c)
        for edge in dominant_component.edges(data=True):
            #print(edge)
            G1.add_edge(edge[0], edge[1], weight=edge[2]['weight'])
        # nx.write_pajek(G, "spectral/spectral3.net")
        nx.write_gml(G1, f"spectral/spectral{k}.gml")
        csizes = [0] * k
        for c in colors:
            csizes[int(c)] += 1
        print(f"Podela na {k}: velicine komponenata su {csizes}")


#make_user_net()
# s_net_t = nx.read_gml('graphs/user_net_with_self_loops.gml')
# print(nx.info(s_net_t))
#read_from_csv_clean_and_save_to_pickle_file()
#read_from_pickle_file()
#print(submissions)
#print(comments)
# print(comments['id'].isnull().values.any())
# print(comments['id'].is_unique)
# print(submissions['subreddit_id'].nunique() + comments['subreddit_id'].nunique())   #1.a)
# grupisano = submissions.groupby(['subreddit', 'subreddit_id'])['num_comments'].agg([('Number of comments per subreddit', np.sum)]).nlargest(10, 'Number of comments per subreddit').reset_index()  #1.c)
read_from_csv(Type.SUBMISSION)
read_from_csv(Type.COMMENT)
_1b()
#_3()
#_4()
#_5()
#_6()
#make_graph()
#G = read_graph_from_file()
#analise_s_net_t(G)
#grapf_laplasijan_s_net_20(G)
#spectral(G)
#print(get_subreddit_id_by_subreddit_name('JailbaitPhotos1'))
# print(submissions.shape[0])
# print(submissions['id'].nunique())
# print(comments.shape[0])
# print(comments['id'].nunique())
# print(comments[comments.duplicated(['id'])])
# print(submissions.dtypes)
# print(comments.dtypes)


#  ako ne postoji subredit dodati novog sa autorom i vremenom
#  ako postoji vec subredit dodati u hashmapu na njegovo mesto jos jednog autora ako taj autor ne postoji i auzirati na najkasnnije vreme
#  proci za objave i komentare
#  proci kroz svaka 2 iz hasmape, ako imaju zajednickog autora stvara se grana

# Proveriti da su submissions i comments sa jedinstvenim i not null id-em
# print(comments['id'].isnull().values.any())
# print(comments['id'].is_unique)