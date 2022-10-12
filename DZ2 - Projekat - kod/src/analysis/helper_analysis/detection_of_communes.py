import math
import numpy as np
import matplotlib.pyplot as plt
import networkx as nx
import pandas as pd
from enum import Enum
from scipy.stats import pearsonr
import seaborn as sns
from matplotlib import pyplot
from selenium import webdriver
from selenium.common.exceptions import WebDriverException, NoSuchElementException
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
import time
from selenium.webdriver.common.by import By
import pickle
from datetime import datetime
from collections import Counter
from operator import itemgetter
import prettytable
from scipy import linalg
from docx import Document
from sklearn.cluster import SpectralClustering
from os.path import exists
from collections import defaultdict
from src.analysis.helper_analysis.basic_analysis import get_giant_component, NetType, read_s_net_giant_component_graph, \
    get_subreddit_name_by_subreddit_id, read_s_net_f_giant_component_graph, read_user_net_giant_component_graph
from src.analysis.statistical_analysis import read_submissions_and_comments

PATH_TO_SPECTRAL_FOLDER = "../../results/spectral"
PATH_TO_SPECTRAL_S_NET_FOLDER = "../../results/spectral/s_net"
PATH_TO_SPECTRAL_S_NET_F_FOLDER = "../../results/spectral/s_net_f"
PATH_TO_SPECTRAL_S_NET_T_FOLDER = "../../results/spectral/s_net_t"
PATH_TO_SPECTRAL_USER_NET_FOLDER = "../../results/spectral/user_net"

S_NET_GIANT_COMPONENT_FILE_PATH = '../../results/nets/giant_component/s_net_giant_component.gml'
S_NET_F_GIANT_COMPONENT_FILE_PATH = '../../results/nets/giant_component/s_net_f_giant_component.gml'
S_NET_T_FILE_PATH = '../../results/nets/s_net_t.gml'
USER_NET_GIANT_COMPONENT_FILE_PATH = '../../../results/nets/giant_component/user_net_giant_component.gml'


def plot_scatter_plot(x_data, y_data, x_label, y_label, title, folder_path):
    fig, ax = plt.subplots()
    ax.scatter(x_data, y_data)
    ax.set_ylabel(y_label, fontsize=15)
    ax.set_xlabel(x_label, fontsize=15)
    ax.set_title(title)
    plt.savefig(f'{ folder_path }/{title}.svg')
    plt.show()


def in_range(value, target, tolerated_error):
    return abs(value-target) < tolerated_error


def get_components3(L, component_number):
    evals, evecs = linalg.eig(L) # each column of evecs is an eigenvector\n",
    # take only real component of each eigenvalue, because eigenvalues of L are guaranteed to be real \n",
    evals = evals.real
    # take only real component of each eigenvector, because eigenvectors of L are guaranteed to be real \n",
    evecs_temp = []
    for e in evecs:
        v = e.real
        evecs_temp.append(v)
    evecs = evecs_temp
    evecs_horizontal_container = np.array(evecs).T #each row is an eigenvector\n",
    zipped = list(zip(evals, evecs_horizontal_container)) # pair eigenvalues and eigenvectors into eigenpairs\n",
    zipped.sort(key=lambda tup: tup[0]) # sort eigenpairs by eigenvalue\n",
    evals, evecs = zip(*zipped) # evecs is now horizontal, but ordered by paired eigenvalue size\n",
    eigenspace0_basis = evecs[:component_number] # take only the first component_number eigenvectors \n",
    #eigenspace0_basis = filter_matrix(eigenspace0_basis, math.pow(10, -15)) # filter values close to 0\n",
    return find_ccs2(eigenspace0_basis, len(eigenspace0_basis[0]))


def find_ccs2(eigenvectors_holder, n):
    components = []
    already_solved = []
    for current in range(0, n):
        if current not in already_solved:
            component = [current]
            for running in range (current+1, n):
                if running not in already_solved:
                    b = True
                    for tracer in range(0, len(eigenvectors_holder)):
                        eigenvector = eigenvectors_holder[tracer]
                        if not in_range(eigenvector[current], eigenvector[running], math.pow(10,-12)):
                            b = False
                    if b:
                        component.append(running)
                        already_solved.append(running)
            already_solved.append(current)
            components.append(component)
    return components


def spectral(graph, folder_path):
    # dominant_component = get_giant_component(graph)
    dominant_component = graph
    max_k = 40
    num_of_nodes = nx.number_of_nodes(graph)
    if num_of_nodes < 40:
        max_k = num_of_nodes
    for k in range(2, max_k):
        clustering = SpectralClustering(n_clusters=k, assign_labels="discretize", affinity="precomputed").fit(nx.adjacency_matrix(dominant_component))
        colors = clustering.labels_
        c_string = []
        for c in colors:
            c_string.append(str(c))
        G1 = nx.Graph()
        for c, label in zip(c_string, dominant_component.nodes()):
            G1.add_node(label, color=c)
        for edge in dominant_component.edges(data=True):
            #print(edge)
            G1.add_edge(edge[0], edge[1], weight=edge[2]['weight'])
        # nx.write_pajek(G, "spectral/spectral3.net")
        nx.write_gml(G1, f"{folder_path}/spectral{k}.gml")
        csizes = [0] * k
        for c in colors:
            csizes[int(c)] += 1
        print(f"Podela na {k}: veličine komponenata su {csizes}")


def _20_graph_laplacian(graph, folder_path):
    L = nx.laplacian_matrix(graph).toarray()
    eigenvalues = linalg.eigvals(L) # TODO: change to .eigenvalsh because we know the matrix is symmetric
    eigenvalues.sort()
    enumerator = np.array(range(1, len(eigenvalues)+1))
    df_eig = pd.DataFrame(list(zip(enumerator, eigenvalues)))
    # write eigenvalue table
    df_eig30 = df_eig[:30]
    df_eig30.columns = ['k', 'lambda_k']
    df_eig30 = df_eig30.astype({'k': 'int32', 'lambda_k': 'float'})
    print(df_eig30)
    doc = Document()
    #doc = docx.Document('tabela_eig.docx')
    t = doc.add_table(df_eig30.shape[0]+1, df_eig30.shape[1])
    t.cell(0, 0).text = r'$k$'
    t.cell(0, 1).text = r'$lambda_k$'
    for i in range(df_eig30.shape[0]):
        for j in range(df_eig30.shape[-1]):
            t.cell(i + 1, j).text = str(df_eig30.values[i, j])
    doc.save(f'{folder_path}/tabela_eig.docx')#PATH_TO_SPECTRAL_FOLDER + '/tabela_eig.docx')
    plot_scatter_plot(enumerator, eigenvalues, r'$k$', r'$λ_k$', 'Ceo spektar graf laplasijana', folder_path)
    # prvih 30 sopstvenih vrednosti,
    # df_eig_30 = df_eig[:30]
    # u snetu krecemo od 800 jer oko tog broja krece nagli skok, videti da se ovde gleda samo gigantska komponenta pa ce biti i manje klastera
    max_k = 40
    num_of_nodes = nx.number_of_nodes(graph)
    if num_of_nodes < 40:
        max_k = num_of_nodes
    df_eig_30 = df_eig[:max_k]
    plot_scatter_plot(df_eig_30.iloc[:, 0], df_eig_30.iloc[:, 1], r'$k$', r'$λ_k$', f'Prvih {max_k} sopstvenih vrednosti graf laplasijana', folder_path)


def detection_of_communes(net_type):
    if net_type == NetType.S_NET:
        graph = nx.read_gml(S_NET_GIANT_COMPONENT_FILE_PATH)
        folder_path = PATH_TO_SPECTRAL_S_NET_FOLDER
    elif net_type == NetType.S_NET_F:
        graph = nx.read_gml(S_NET_F_GIANT_COMPONENT_FILE_PATH)
        folder_path = PATH_TO_SPECTRAL_S_NET_F_FOLDER
    elif net_type == NetType.S_NET_T:
        graph = nx.read_gml(S_NET_T_FILE_PATH)
        folder_path = PATH_TO_SPECTRAL_S_NET_T_FOLDER
    elif net_type == NetType.USER_NET:
        print('Funkcija nx.laplacian_matrix(graph) ne može da se korisit kod usmerenih grafova, stoga bi bilo bolje korisiti neku drugu metodu.')
        return
    else:
        return
    _20_graph_laplacian(graph, folder_path)
    spectral(graph, folder_path)


#detection_of_communes(NetType.USER_NET)
    # submissions, comments = read_submissions_and_comments()
    # s = ['t5_2qok8', 't5_2qkot', 't5_2qoxq']
    # for i in range(0, len(s)):
    #     print(get_subreddit_name_by_subreddit_id(s[i]))
    #     s1 = submissions.loc[submissions['subreddit_id'] == s[i]]
    #     print(s1)
    #     # for index, row in s1:
    #     #     print(index, row)
    #     # c1 = comments.loc[comments['subreddit_id'] == s[i]]
    #     # for index, row in c1:
    #     #     print(row['author'], datetime.fromtimestamp(row['created_utc']))
    #
    #     # print(submissions.loc[submissions['subreddit_id'] == s[i]]['created_utc'])
    #     # print(comments.loc[comments['subreddit_id'] == s[i]]['created_utc'])
    #     print('-------------------------')
    # # a = ['Cetatenie', 'dezvoltarepersonala', 'positivethinker']